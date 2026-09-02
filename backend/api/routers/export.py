"""Export router: generate styled PDF, Word, CSV and Markdown exports.

The caller posts either a single exchange (``query`` + ``answer``, the JSON
shape returned by ``POST /api/v1/chat``) or a whole conversation (``items``,
a list of ``{query, answer}`` exchanges), so exports work for both a single
response and a full chat session without re-running the graph.
"""

from __future__ import annotations

import csv
import io
import re
from datetime import datetime
from typing import Any, Callable

from fastapi import APIRouter, Depends, HTTPException, Request
from fastapi.responses import Response
from pydantic import BaseModel, Field

from backend.api.deps import get_ctx, get_current_user
from backend.core.models import Citation, FinalAnswer
from backend.security.jwt import TokenPayload

router = APIRouter(tags=["export"])


class ExportItem(BaseModel):
    """One question/answer exchange of a conversation."""

    query: str = Field(default="", description="Original user question")
    answer: FinalAnswer


class ExportRequest(BaseModel):
    """Payload accepted by the export endpoints.

    Either ``items`` (full or partial conversation) or the legacy single
    ``query``/``answer`` pair must be provided.
    """

    query: str = Field(default="", description="Original user question")
    answer: FinalAnswer | None = None
    items: list[ExportItem] | None = Field(
        default=None, description="Conversation exchanges to export"
    )
    session_id: str = Field(default="", description="Conversation session id")
    latency_ms: float = Field(default=0.0)

    def exchanges(self) -> list[ExportItem]:
        if self.items:
            return self.items
        if self.answer is not None:
            return [ExportItem(query=self.query, answer=self.answer)]
        raise ValueError("Export payload must contain 'items' or 'answer'.")


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _strip_markdown(text: str) -> str:
    """Best-effort markdown to plain text for Word/PDF exports."""
    text = re.sub(r"```.*?```", lambda m: m.group(0).strip("`"), text, flags=re.DOTALL)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"__([^_]+)__", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\(([^)]+)\)", r"\1 (\2)", text)
    text = re.sub(r"^#{1,6}\s+", "", text, flags=re.MULTILINE)
    text = text.replace("•", "-")
    return text.strip()


def _citation_ref(citation: Citation, index: int) -> str:
    """Reference marker for a citation.

    Labels produced by the pipeline already carry the marker used inside the
    answer text (e.g. ``[3]``); reuse it so the export matches the answer and
    never shows doubled numbers like ``[1] [3]``.
    """
    label = (citation.label or "").strip()
    return label if re.fullmatch(r"\[\d+\]", label) else f"[{index}]"


def _merged_citations(citations: list[Citation]) -> list[tuple[str, Citation]]:
    """Merge citations that point to the same document and article.

    Retrieval often yields several chunks of the same article, producing
    repeated entries like ``[9] — Charte…, art. 168`` / ``[10] — Charte…,
    art. 168``. They are merged into a single line whose markers are joined
    (``[9], [10] — …``) so every marker used in the answer still resolves.
    """
    merged: list[tuple[list[str], Citation]] = []
    index: dict[tuple[str, str], int] = {}
    for i, citation in enumerate(citations, 1):
        ref = _citation_ref(citation, i)
        key = (citation.document_name, citation.article or "")
        if key in index:
            merged[index[key]][0].append(ref)
        else:
            index[key] = len(merged)
            merged.append(([ref], citation))
    return [(", ".join(refs), citation) for refs, citation in merged]


def _citation_label(citation: Citation, ref: str) -> str:
    """Descriptive label when it is not already the bare ``[n]`` marker."""
    label = _strip_markdown(citation.label or "").strip()
    return "" if not label or re.fullmatch(r"\[\d+\]", label) else label


# ---------------------------------------------------------------------------
# PDF export
# ---------------------------------------------------------------------------


def _build_pdf(data: ExportRequest, title: str = "Réponse juridique") -> bytes:
    from reportlab.lib import colors
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.lib.units import cm
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer

    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        rightMargin=2 * cm,
        leftMargin=2 * cm,
        topMargin=2 * cm,
        bottomMargin=2 * cm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "TitleCustom",
        parent=styles["Title"],
        fontSize=18,
        textColor=colors.HexColor("#1e3a8a"),
        spaceAfter=14,
        alignment=1,
    )
    heading_style = ParagraphStyle(
        "HeadingCustom",
        parent=styles["Heading2"],
        fontSize=13,
        textColor=colors.HexColor("#1e40af"),
        spaceBefore=14,
        spaceAfter=8,
    )
    normal_style = ParagraphStyle(
        "NormalCustom",
        parent=styles["BodyText"],
        fontSize=10,
        leading=14,
        spaceAfter=6,
    )
    small_style = ParagraphStyle(
        "SmallCustom",
        parent=styles["BodyText"],
        fontSize=9,
        leading=12,
        textColor=colors.HexColor("#374151"),
    )
    question_style = ParagraphStyle(
        "QuestionCustom",
        parent=normal_style,
        textColor=colors.HexColor("#1e3a8a"),
        spaceBefore=10,
    )

    items = data.exchanges()
    multiple = len(items) > 1

    story: list[Any] = []
    story.append(Paragraph(title, title_style))

    for n, item in enumerate(items, 1):
        answer = item.answer
        if multiple:
            story.append(Paragraph(f"Échange {n}", heading_style))
        if item.query:
            story.append(
                Paragraph(f"<b>Question :</b> {_strip_markdown(item.query)}", question_style)
            )

        if answer.requires_human_review:
            story.append(Spacer(1, 8))
            story.append(
                Paragraph(
                    "<b>Révision humaine requise</b> — cette réponse doit être validée par un juriste.",
                    ParagraphStyle(
                        "Warning",
                        parent=normal_style,
                        textColor=colors.HexColor("#991b1b"),
                        backColor=colors.HexColor("#fee2e2"),
                        borderPadding=6,
                    ),
                )
            )

        story.append(Spacer(1, 12))
        story.append(Paragraph("Réponse", heading_style))
        for paragraph in _strip_markdown(answer.answer).split("\n\n"):
            if paragraph.strip():
                story.append(Paragraph(paragraph.replace("\n", "<br/>"), normal_style))

        if answer.citations:
            story.append(Paragraph("Citations", heading_style))
            for ref, citation in _merged_citations(answer.citations):
                label = _citation_label(citation, ref)
                line = f"{ref} "
                line += f"<b>{label}</b> — " if label else "— "
                line += citation.document_name
                if citation.article:
                    line += f", art. {citation.article}"
                status = "✓ vérifiée" if citation.verified else "non vérifiée"
                line += f" <i>({status})</i>"
                if citation.url:
                    line += f"<br/><a href='{citation.url}' color='blue'>{citation.url}</a>"
                story.append(Paragraph(line, small_style))
                story.append(Spacer(1, 4))

        if answer.warnings:
            story.append(Paragraph("Avertissements", heading_style))
            for warning in answer.warnings:
                story.append(Paragraph(f"• {_strip_markdown(warning)}", small_style))

    story.append(Spacer(1, 20))
    story.append(
        Paragraph(
            "Avertissement : cette réponse est une aide à la recherche juridique. "
            "Elle ne constitue pas un conseil juridique. Consultez un professionnel du droit.",
            ParagraphStyle(
                "Disclaimer",
                parent=small_style,
                textColor=colors.HexColor("#92400e"),
                alignment=1,
            ),
        )
    )
    story.append(
        Paragraph(
            f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} — Assistant Juridique Burkina Faso",
            ParagraphStyle("Footer", parent=small_style, alignment=1, fontSize=8),
        )
    )

    doc.build(story)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# Word export
# ---------------------------------------------------------------------------


def _build_docx(data: ExportRequest, title: str = "Réponse juridique") -> bytes:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches, Pt, RGBColor

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.8)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

    title_heading = doc.add_heading(title, level=0)
    title_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_heading.runs[0]
    run.font.color.rgb = RGBColor(30, 58, 138)
    run.font.size = Pt(20)

    items = data.exchanges()
    multiple = len(items) > 1

    for n, item in enumerate(items, 1):
        answer = item.answer
        if multiple:
            doc.add_heading(f"Échange {n}", level=1)
        if item.query:
            p = doc.add_paragraph()
            p.add_run("Question : ").bold = True
            p.add_run(_strip_markdown(item.query))

        if answer.requires_human_review:
            p = doc.add_paragraph()
            p.add_run("Révision humaine requise").bold = True
            p.add_run(" — cette réponse doit être validée par un juriste.")
            p.runs[0].font.color.rgb = RGBColor(153, 27, 27)

        doc.add_heading("Réponse", level=2 if multiple else 1)
        for paragraph in _strip_markdown(answer.answer).split("\n\n"):
            if paragraph.strip():
                doc.add_paragraph(paragraph)

        if answer.citations:
            doc.add_heading("Citations", level=2 if multiple else 1)
            for ref, citation in _merged_citations(answer.citations):
                label = _citation_label(citation, ref)
                p = doc.add_paragraph()
                p.add_run(f"{ref} ").bold = True
                p.add_run(f"{label} — " if label else "— ")
                p.add_run(citation.document_name)
                if citation.article:
                    p.add_run(f", art. {citation.article}")
                status = "vérifiée" if citation.verified else "non vérifiée"
                p.add_run(f" ({status})")
                if citation.url:
                    p.add_run(f" — {citation.url}")

        if answer.warnings:
            doc.add_heading("Avertissements", level=2 if multiple else 1)
            for warning in answer.warnings:
                doc.add_paragraph(warning, style="List Bullet")

    disclaimer = doc.add_paragraph()
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = disclaimer.add_run(
        "Avertissement : cette réponse est une aide à la recherche juridique. "
        "Elle ne constitue pas un conseil juridique. Consultez un professionnel du droit."
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(146, 64, 14)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run(
        f"Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} — Assistant Juridique Burkina Faso"
    )
    footer_run.font.size = Pt(8)

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()


# ---------------------------------------------------------------------------
# CSV export
# ---------------------------------------------------------------------------


def _build_csv(data: ExportRequest, title: str = "Réponse juridique") -> bytes:
    # ``title`` is accepted for a uniform builder signature; the CSV layout
    # keeps its own long-standing header instead.
    buffer = io.StringIO()
    writer = csv.writer(buffer)
    writer.writerow(["Assistant Juridique Burkina Faso — Export CSV"])
    for n, item in enumerate(data.exchanges(), 1):
        answer = item.answer
        writer.writerow([])
        writer.writerow([f"Échange {n}"])
        writer.writerow(["Question", _strip_markdown(item.query)])
        writer.writerow(["Révision humaine", "Oui" if answer.requires_human_review else "Non"])
        writer.writerow(["Réponse"])
        for paragraph in _strip_markdown(answer.answer).split("\n\n"):
            writer.writerow([paragraph])
        writer.writerow(["Citations"])
        writer.writerow(["#", "Label", "Document", "Article", "Vérifiée", "URL"])
        for ref, citation in _merged_citations(answer.citations):
            writer.writerow(
                [
                    ref,
                    _citation_label(citation, ref),
                    citation.document_name,
                    citation.article or "",
                    "Oui" if citation.verified else "Non",
                    citation.url or "",
                ]
            )
        if answer.warnings:
            writer.writerow(["Avertissements"])
            for warning in answer.warnings:
                writer.writerow([warning])
    writer.writerow([])
    writer.writerow(
        [
            "Avertissement : cette réponse est une aide à la recherche juridique. "
            "Elle ne constitue pas un conseil juridique. Consultez un professionnel du droit."
        ]
    )
    return buffer.getvalue().encode("utf-8-sig")


# ---------------------------------------------------------------------------
# Markdown export
# ---------------------------------------------------------------------------


def _build_markdown(data: ExportRequest, title: str = "Réponse juridique") -> bytes:
    """Clean French markdown document; the answer body keeps its markdown."""
    lines: list[str] = [
        f"# {title}",
        "",
        f"*Généré le {datetime.now().strftime('%d/%m/%Y à %H:%M')} — Assistant Juridique Burkina Faso*",
        "",
    ]
    items = data.exchanges()
    multiple = len(items) > 1

    for n, item in enumerate(items, 1):
        answer = item.answer
        if multiple:
            lines += [f"## Échange {n}", ""]
        if item.query:
            lines += [f"**Question :** {item.query}", ""]
        if answer.requires_human_review:
            lines += [
                "> **Révision humaine requise** — cette réponse doit être validée par un juriste.",
                "",
            ]
        lines += ["### Réponse" if multiple else "## Réponse", "", answer.answer.strip(), ""]

        verified = [c for c in answer.citations if c.verified]
        if verified:
            lines += ["### Références" if multiple else "## Références", ""]
            for ref, citation in _merged_citations(verified):
                label = _citation_label(citation, ref)
                entry = f"- **{ref}**"
                if label:
                    entry += f" **{label}**"
                if citation.article:
                    entry += f", art. {citation.article}"
                if citation.document_name:
                    entry += f" — {citation.document_name}"
                if citation.url:
                    entry += f" — {citation.url}"
                lines.append(entry)
            lines.append("")

        if answer.warnings:
            lines += ["### Avertissements" if multiple else "## Avertissements", ""]
            lines += [f"- {_strip_markdown(w)}" for w in answer.warnings]
            lines.append("")

    lines += [
        "---",
        "",
        "_Avertissement : cette réponse est une aide à la recherche juridique. "
        "Elle ne constitue pas un conseil juridique. Consultez un professionnel du droit._",
        "",
    ]
    return "\n".join(lines).encode("utf-8")


# ---------------------------------------------------------------------------
# Router
# ---------------------------------------------------------------------------


_FORMATS: dict[str, tuple[str, str, Callable[[ExportRequest], bytes]]] = {
    "pdf": ("application/pdf", ".pdf", _build_pdf),
    "word": (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        ".docx",
        _build_docx,
    ),
    "csv": ("text/csv; charset=utf-8", ".csv", _build_csv),
    "md": ("text/markdown; charset=utf-8", ".md", _build_markdown),
}


@router.post("/export/{format}")
async def export_answer(
    request: Request,
    format: str,
    payload: ExportRequest,
    # Auth: generation is CPU-heavy, so the endpoint is no longer anonymous.
    # get_current_user (not require_user) keeps the development anonymous
    # pass-through; production rejects missing/invalid tokens with a 401.
    _user: TokenPayload = Depends(get_current_user),
) -> Response:
    """Export one answer or a whole conversation to PDF, Word, CSV or Markdown."""
    format = format.lower()
    if format not in _FORMATS:
        raise HTTPException(
            status_code=400,
            detail=f"Unsupported format: {format}. Supported: {', '.join(_FORMATS)}.",
        )
    ctx = get_ctx(request)
    mime, ext, builder = _FORMATS[format]
    prefix = "conversation-juridique" if payload.items and len(payload.items) > 1 else "reponse-juridique"
    filename = f"{prefix}-{datetime.now().strftime('%Y%m%d-%H%M%S')}{ext}"
    try:
        data = builder(payload, title=ctx.settings.export_pdf_title)
    except Exception as exc:
        raise HTTPException(status_code=500, detail=f"Export failed: {exc}") from exc
    return Response(
        content=data,
        media_type=mime,
        headers={"Content-Disposition": f"attachment; filename=\"{filename}\""},
    )
