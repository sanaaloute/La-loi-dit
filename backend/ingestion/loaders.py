"""Document loaders: PDF, DOCX, HTML, plain text, Markdown and CSV.

All third-party imports (PyMuPDF, python-docx, BeautifulSoup, Pillow, httpx)
happen lazily inside the loader functions so this module imports cleanly even
when optional dependencies are missing. Loaders raise
:class:`backend.core.exceptions.IngestionError` on unreadable files.
"""

from __future__ import annotations

import asyncio
import csv
import io
import logging
import re
import shutil
import tempfile
from pathlib import Path
from typing import Any, Optional, Union

from pydantic import BaseModel, Field

from backend.core.exceptions import IngestionError

logger = logging.getLogger(__name__)

_URL_RE = re.compile(r"^https?://", re.IGNORECASE)

#: Pages with less extracted text than this are treated as scanned and OCR'd.
_OCR_TEXT_THRESHOLD = 20

#: Extensions dispatched by :func:`load_any`.
SUPPORTED_EXTENSIONS = {
    ".pdf": "pdf",
    ".docx": "docx",
    ".html": "html",
    ".htm": "html",
    ".txt": "txt",
    ".md": "markdown",
    ".markdown": "markdown",
    ".csv": "csv",
}


class ExtractedDocument(BaseModel):
    """Raw text extracted from a source document, before cleaning/chunking."""

    name: str
    text: str = ""
    pages: list[str] = Field(default_factory=list)
    metadata: dict[str, Any] = Field(default_factory=dict)


def _read_text_file(path: Path) -> str:
    """Read a text file, tolerating common encodings (UTF-8 first)."""
    raw = path.read_bytes()
    for encoding in ("utf-8", "utf-8-sig", "cp1252", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="replace")


def load_pdf(path: Union[str, Path]) -> ExtractedDocument:
    """Extract per-page text from a PDF via PyMuPDF.

    Pages yielding little or no text (scanned pages) are recorded in
    ``metadata["ocr_needed_pages"]``; when OCR is available and enabled they
    are re-rendered at 300 dpi and recognized with PaddleOCR (best-effort,
    never fatal), with the outcome noted in ``metadata["ocr_recovered_pages"]``
    / ``metadata["ocr_skipped_pages"]``.
    """
    p = Path(path)
    if not p.is_file():
        raise IngestionError(f"PDF not found: {p}")
    try:
        import pymupdf
    except ImportError:
        try:
            import fitz as pymupdf  # older PyMuPDF exposes only the fitz module
        except ImportError as exc:
            raise IngestionError("PyMuPDF is required to load PDF files") from exc

    try:
        with pymupdf.open(str(p)) as pdf:
            pages: list[str] = []
            for page in pdf:
                try:
                    pages.append((page.get_text() or "").strip())
                except Exception:
                    pages.append("")  # a broken page must not kill the document

            metadata: dict[str, Any] = {"loader": "pdf", "page_count": len(pages), "path": str(p)}
            ocr_needed = [i + 1 for i, text in enumerate(pages) if len(text) < _OCR_TEXT_THRESHOLD]
            if ocr_needed:
                metadata["ocr_needed_pages"] = ocr_needed
                _ocr_scanned_pages(pdf, pages, ocr_needed, metadata)
    except IngestionError:
        raise
    except Exception as exc:
        raise IngestionError(f"Failed to read PDF {p}: {exc}") from exc

    return ExtractedDocument(name=p.name, text="\n\n".join(pages), pages=pages, metadata=metadata)


def _ocr_scanned_pages(pdf: Any, pages: list[str], ocr_needed: list[int], metadata: dict[str, Any]) -> None:
    """OCR textless pages with PaddleOCR. Best-effort, never raises."""
    from backend.ingestion import ocr as ocr_engine

    metadata["ocr_available"] = ocr_engine.ocr_available()
    if not metadata["ocr_available"]:
        return
    from backend.core.config import get_settings

    max_pages = get_settings().ocr_max_pages
    to_process, skipped = ocr_needed[:max_pages], ocr_needed[max_pages:]
    if skipped:
        logger.warning("OCR page cap (%d) reached; skipping pages %s", max_pages, skipped)
        metadata["ocr_skipped_pages"] = skipped
    # Render the pages to PNG files, then OCR the whole batch in a single
    # child process (paddle never loads in this process — see ingestion/ocr.py).
    tmp_dir = Path(tempfile.mkdtemp(prefix="ocr_pages_"))
    try:
        rendered: list[tuple[int, Path]] = []
        for page_no in to_process:
            try:
                image_path = tmp_dir / f"page_{page_no}.png"
                image_path.write_bytes(pdf[page_no - 1].get_pixmap(dpi=300).tobytes("png"))
                rendered.append((page_no, image_path))
            except Exception:
                continue  # a broken page must not kill the document
        results = ocr_engine.ocr_images([image_path for _, image_path in rendered])
        recovered: list[int] = []
        for page_no, image_path in rendered:
            text = results.get(str(image_path), "").strip()
            if len(text) > len(pages[page_no - 1]):
                pages[page_no - 1] = text
                recovered.append(page_no)
        if recovered:
            metadata["ocr_recovered_pages"] = recovered
    finally:
        shutil.rmtree(tmp_dir, ignore_errors=True)


def load_docx(path: Union[str, Path]) -> ExtractedDocument:
    """Extract paragraphs and table cells from a DOCX via python-docx."""
    p = Path(path)
    if not p.is_file():
        raise IngestionError(f"DOCX not found: {p}")
    try:
        import docx
    except ImportError as exc:
        raise IngestionError("python-docx is required to load DOCX files") from exc

    try:
        document = docx.Document(str(p))
        parts = [para.text for para in document.paragraphs if para.text.strip()]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    parts.append(line)
    except Exception as exc:
        raise IngestionError(f"Failed to read DOCX {p}: {exc}") from exc

    text = "\n".join(parts)
    return ExtractedDocument(
        name=p.name,
        text=text,
        pages=[text],  # DOCX has no stable pagination
        metadata={"loader": "docx", "path": str(p)},
    )


def load_html(path_or_url: Union[str, Path], content: Optional[str] = None) -> ExtractedDocument:
    """Extract main text from HTML via BeautifulSoup.

    ``content`` may be supplied directly (e.g. by the crawler). Otherwise a
    local file is read, or an ``http(s)://`` URL is fetched with httpx
    (lazy import); network failures raise :class:`IngestionError`.
    """
    try:
        from bs4 import BeautifulSoup
    except ImportError as exc:
        raise IngestionError("beautifulsoup4 is required to load HTML") from exc

    source = str(path_or_url)
    metadata: dict[str, Any] = {"loader": "html"}
    if content is None:
        if _URL_RE.match(source):
            metadata["url"] = source
            try:
                import httpx
            except ImportError as exc:
                raise IngestionError("httpx is required to fetch HTML URLs") from exc
            try:
                from backend.core.config import get_settings

                timeout = get_settings().ingestion_html_timeout_seconds
                response = httpx.get(source, timeout=timeout, follow_redirects=True)
                response.raise_for_status()
                content = response.text
            except Exception as exc:
                raise IngestionError(f"Failed to fetch {source}: {exc}") from exc
            name = source.rstrip("/").rsplit("/", 1)[-1] or source
        else:
            p = Path(source)
            if not p.is_file():
                raise IngestionError(f"HTML file not found: {p}")
            content = _read_text_file(p)
            name = p.name
            metadata["path"] = str(p)
    else:
        name = Path(source).name if not _URL_RE.match(source) else source
        if _URL_RE.match(source):
            metadata["url"] = source
        else:
            metadata["path"] = source

    try:
        try:
            soup = BeautifulSoup(content, "lxml")
        except Exception:
            soup = BeautifulSoup(content, "html.parser")
        for tag in soup(["script", "style", "noscript", "header", "footer", "nav"]):
            tag.decompose()
        title = soup.title.string.strip() if soup.title and soup.title.string else ""
        if title:
            metadata["title"] = title
        text = soup.get_text("\n")
    except Exception as exc:
        raise IngestionError(f"Failed to parse HTML from {source}: {exc}") from exc

    return ExtractedDocument(name=name, text=text, pages=[text], metadata=metadata)


def load_txt(path: Union[str, Path]) -> ExtractedDocument:
    """Load a plain-text file."""
    p = Path(path)
    if not p.is_file():
        raise IngestionError(f"Text file not found: {p}")
    try:
        text = _read_text_file(p)
    except Exception as exc:
        raise IngestionError(f"Failed to read {p}: {exc}") from exc
    return ExtractedDocument(name=p.name, text=text, pages=[text], metadata={"loader": "txt", "path": str(p)})


def load_markdown(path: Union[str, Path]) -> ExtractedDocument:
    """Load a Markdown file, stripping the most common markup markers."""
    p = Path(path)
    if not p.is_file():
        raise IngestionError(f"Markdown file not found: {p}")
    try:
        raw = _read_text_file(p)
    except Exception as exc:
        raise IngestionError(f"Failed to read {p}: {exc}") from exc
    text = _strip_markdown(raw)
    return ExtractedDocument(
        name=p.name,
        text=text,
        pages=[text],
        metadata={"loader": "markdown", "path": str(p)},
    )


def load_csv(path: Union[str, Path]) -> ExtractedDocument:
    """Load a CSV file, rendering rows as text lines with the header preserved.

    Uses the stdlib ``csv`` module with the same encoding tolerance as
    :func:`load_txt`. Each row becomes a ``header=value`` line so the cells
    stay self-describing once chunked out of order.
    """
    p = Path(path)
    if not p.is_file():
        raise IngestionError(f"CSV file not found: {p}")
    try:
        raw = _read_text_file(p)
    except Exception as exc:
        raise IngestionError(f"Failed to read {p}: {exc}") from exc
    try:
        rows = list(csv.reader(io.StringIO(raw)))
    except csv.Error as exc:
        raise IngestionError(f"Failed to parse CSV {p}: {exc}") from exc
    if not rows:
        raise IngestionError(f"Empty CSV file: {p}")

    header = rows[0]
    lines = [" | ".join(header)]
    for row in rows[1:]:
        if not any(cell.strip() for cell in row):
            continue  # skip blank lines
        cells = [
            f"{header[i]}={cell}" if i < len(header) and header[i].strip() else cell
            for i, cell in enumerate(row)
        ]
        lines.append(" | ".join(cells))
    text = "\n".join(lines)
    return ExtractedDocument(
        name=p.name,
        text=text,
        pages=[text],
        metadata={"loader": "csv", "format": "csv", "path": str(p), "row_count": len(lines) - 1},
    )


def _strip_markdown(raw: str) -> str:
    """Lightweight Markdown-to-text (headings, links, emphasis, code fences)."""
    text = re.sub(r"```.*?```", lambda m: m.group(0).strip("`"), raw, flags=re.DOTALL)
    text = re.sub(r"!\[([^\]]*)\]\([^)]*\)", r"\1", text)  # images -> alt text
    text = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", text)  # links -> anchor text
    text = re.sub(r"^#{1,6}\s*", "", text, flags=re.MULTILINE)  # heading markers
    text = re.sub(r"(\*\*|__)(.*?)\1", r"\2", text)
    text = re.sub(r"(\*|_)(.*?)\1", r"\2", text)
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)  # bullets
    return text


async def load_any(path: Union[str, Path]) -> ExtractedDocument:
    """Dispatch to the right loader based on file extension.

    Raises :class:`IngestionError` for unsupported extensions or unreadable
    files. Runs the synchronous loaders in a worker thread.
    """
    p = Path(path)
    kind = SUPPORTED_EXTENSIONS.get(p.suffix.lower())
    if kind is None:
        raise IngestionError(
            f"Unsupported file extension '{p.suffix}' for {p}. "
            f"Supported: {sorted(SUPPORTED_EXTENSIONS)}"
        )
    loaders = {
        "pdf": load_pdf,
        "docx": load_docx,
        "html": load_html,
        "txt": load_txt,
        "markdown": load_markdown,
        "csv": load_csv,
    }
    return await asyncio.to_thread(loaders[kind], p)
